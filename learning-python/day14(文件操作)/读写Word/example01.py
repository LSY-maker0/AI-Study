"""
example01 - Python读写Word文件

安装第三方库 python-docx
pillow(处理图片) 操作图片，关于图片的各种信息
pillow -> PIL -> Python Image Library

Author: lsy
Date: 2026/1/6
"""
from docx import Document
from docx.shared import Inches

from docx.document import Document as Doc

document = Document() # type: Doc # 1. 创建word文档对象（构造器语法创建对象）

document.add_heading('快乐学Python，大标题', 0) # 添加大标题 0

p = document.add_paragraph('Python是一门流行的编程语言，它')
p.add_run('简单').bold = True
p.add_run('而且')
p.add_run('优雅').italic = True

document.add_heading('一级标题', level=1)
document.add_paragraph('1. 段落 Intense Quote样式', style='Intense Quote')

document.add_paragraph(
    '2. 段落', style='List Bullet'
)
document.add_paragraph(
    '3. 段落', style='List Number'
)

document.add_picture('../resources/apple.jpeg', width=Inches(1.25))

records = (
    ('张三', '男', 'Spam'),
    ('李四', '男', 'Eggs'),
    ('王五', '女', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
hdr_cells[1].text = '性别'
hdr_cells[2].text = '3'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('../resources/demo.docx')
